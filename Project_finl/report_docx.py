from __future__ import annotations

from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

# Paths and metadata
ROOT = Path(__file__).resolve().parent
OUTDIR = ROOT / "outputs"
DOCX_PATH = OUTDIR / "report.docx"
DATA_PATH = ROOT / "email-Eu-core.txt"

STUDENT_NAME = "Priyansh Kumar Paswan"
ROLL_NO = "205124071"
SUBJECT = "Social Network Analysis"
PROFESSOR = "Dr. S.R. Balasundaram"
DEPARTMENT = "Computer Applications"
INSTITUTE = "National Institute of Technology, Tiruchirappalli"
LOGO_PATH = Path("/Users/prx./Documents/SNA/Project_finl/National_Institute_of_Technology,_Tiruchirappalli.svg.png")


def ensure_outputs_exist():
    needed = [
        OUTDIR / "subgraph.png",
        OUTDIR / "degree_distribution.png",
        OUTDIR / "link_analysis.csv",
        OUTDIR / "link_analysis_top_pagerank.png",
        OUTDIR / "influence_scores.csv",
        OUTDIR / "influence_top.png",
        OUTDIR / "node_classification_labels.csv",
        OUTDIR / "link_prediction_roc.png",
        OUTDIR / "link_prediction_auc.csv",
        OUTDIR / "anomaly_scatter.png",
        OUTDIR / "anomalies.csv",
        OUTDIR / "graph_attributes.gexf",
    ]
    missing = [p for p in needed if not p.exists()]
    if missing:
        import subprocess, sys
        print("Some outputs missing; running project.py to generate them...")
        cmd = [sys.executable, str(ROOT / "project.py"), "--task", "all", "--outdir", str(OUTDIR)]
        subprocess.check_call(cmd, cwd=ROOT)


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_table(doc: Document, df: pd.DataFrame, title: str = None, max_rows: int = 15, cols=None):
    if title:
        add_heading(doc, title, level=3)
    if cols:
        df = df[cols]
    df = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr_cells[i].text = str(c)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(df.columns):
            cells[i].text = str(row[c])


def add_image(doc: Document, path: Path, caption: str, width_inches: float = 6.0):
    if path.exists():
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width_inches))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_placeholder(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"** {text} **")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def build_docx():
    ensure_outputs_exist()

    doc = Document()
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    # Title page
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(LOGO_PATH), width=Inches(2.6))
    t = doc.add_paragraph(SUBJECT)
    t.style = doc.styles['Title']
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Project Report")
    sub.style = doc.styles['Heading 2']
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    inst = doc.add_paragraph(INSTITUTE)
    inst.style = doc.styles['Heading 3']
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    add_table(doc, pd.DataFrame([
        ["Student", STUDENT_NAME],
        ["Roll Number", ROLL_NO],
        ["Professor", PROFESSOR],
        ["Department", DEPARTMENT],
        ["Institute", INSTITUTE],
    ], columns=["Field", "Value"]), title=None, max_rows=5)

    doc.add_page_break()

    # 1. Overview
    add_heading(doc, "1. Dataset Overview", level=2)
    add_paragraph(doc, "Email-EU-core undirected graph. The analysis below summarizes structure and degree characteristics.")

    # Basic stats table
    import networkx as nx
    G = nx.read_edgelist(DATA_PATH, create_using=nx.Graph(), nodetype=int)
    nodes = G.number_of_nodes(); edges = G.number_of_edges()
    density = nx.density(G); clustering = nx.average_clustering(G); is_conn = nx.is_connected(G)
    stats_df = pd.DataFrame([
        ["Nodes", nodes],
        ["Edges", edges],
        ["Density", f"{density:.4f}"],
        ["Avg clustering", f"{clustering:.4f}"],
        ["Connected", f"{is_conn}"],
    ], columns=["Metric", "Value"])
    add_table(doc, stats_df, title=None, max_rows=10)

    add_image(doc, OUTDIR / "subgraph.png", "Figure 1: Subgraph visualization")
    add_image(doc, OUTDIR / "degree_distribution.png", "Figure 2: Degree distribution histogram")
    add_placeholder(doc, "Optional Gephi screenshot: Global layout (ForceAtlas2), color by community, size by pagerank")

    doc.add_page_break()

    # 2. Link Analysis
    add_heading(doc, "2. Link Analysis (PageRank & Eigenvector)", level=2)
    add_paragraph(doc, "Methods: PageRank for importance via random walks; Eigenvector centrality for importance via influential neighbors.")
    la = pd.read_csv(OUTDIR / "link_analysis.csv").sort_values("pagerank", ascending=False)
    add_table(doc, la, title="Top-10 by PageRank", max_rows=10, cols=["node", "pagerank", "eigenvector", "degree"])
    add_image(doc, OUTDIR / "link_analysis_top_pagerank.png", "Figure 3: Top nodes by PageRank (bar chart)")
    add_placeholder(doc, "Optional Gephi screenshot: Influence view — size by pagerank, color gradient by eigenvector")

    doc.add_page_break()

    # 3. Node Classification
    add_heading(doc, "3. Node Classification (Label Propagation)", level=2)
    add_paragraph(doc, "Method: Asynchronous label propagation to discover communities without ground truth labels.")
    labels = pd.read_csv(OUTDIR / "node_classification_labels.csv")
    comm_counts = labels.groupby("community").size().reset_index(name="size").sort_values("size", ascending=False)
    add_table(doc, comm_counts, title="Communities by size (Top 15)", max_rows=15, cols=["community", "size"])
    add_placeholder(doc, "ADD GEPHI SCREENSHOT HERE: Community visualization — color by community (categorical), size by pagerank; layout: ForceAtlas2")

    doc.add_page_break()

    # 4. Influence Analysis
    add_heading(doc, "4. Influence Analysis (PageRank & Betweenness)", level=2)
    add_paragraph(doc, "Methods: PageRank for global influence; Betweenness for brokerage across communities.")
    inf = pd.read_csv(OUTDIR / "influence_scores.csv").sort_values(["pagerank", "betweenness"], ascending=False)
    add_table(doc, inf, title="Top-10 influencers", max_rows=10, cols=["node", "pagerank", "betweenness", "degree"])
    add_image(doc, OUTDIR / "influence_top.png", "Figure 4: Top influencers (PageRank + Betweenness)")
    add_placeholder(doc, "ADD GEPHI SCREENSHOT HERE: Influencers — size by betweenness or pagerank; optionally label top nodes")

    doc.add_page_break()

    # 5. Link Prediction
    add_heading(doc, "5. Link Prediction (Adamic-Adar, Jaccard, Preferential Attachment)", level=2)
    add_paragraph(doc, "Procedure: Hold out 10% edges; score with three heuristics on the train graph; evaluate ROC/AUC.")
    auc = pd.read_csv(OUTDIR / "link_prediction_auc.csv").sort_values("auc", ascending=False)
    add_table(doc, auc, title="AUC by metric", max_rows=5, cols=["metric", "auc"])
    add_image(doc, OUTDIR / "link_prediction_roc.png", "Figure 5: ROC curves for link prediction metrics")

    doc.add_page_break()

    # 6. Anomaly Detection
    add_heading(doc, "6. Anomaly Detection (IsolationForest on egonet features)", level=2)
    add_paragraph(doc, "Features: degree, clustering, average neighbor degree, egonet edges; Model: IsolationForest (1% contamination).")
    anom = pd.read_csv(OUTDIR / "anomalies.csv")
    anom_top = anom.sort_values("decision_function").head(10)
    add_table(doc, anom_top, title="Top-10 anomalous nodes", max_rows=10, cols=["node", "degree", "clustering", "avg_neighbor_degree", "ego_edges", "decision_function"])
    add_image(doc, OUTDIR / "anomaly_scatter.png", "Figure 6: Anomaly scatter (Degree vs Clustering, colored by score)")
    add_placeholder(doc, "ADD GEPHI SCREENSHOT HERE: Anomaly view — color by anomaly_score (continuous), highlight top outliers")

    doc.add_page_break()

    # 7. Gephi Visualizations
    add_heading(doc, "7. Gephi Visualizations (What to include)", level=2)
    add_paragraph(doc, "Open outputs/graph_attributes.gexf in Gephi and capture these views:")
    for b in [
        "Community view: Color by 'community' (categorical), size by 'pagerank'.",
        "Influence view: Size by 'betweenness' or 'pagerank'; optional color gradient by betweenness.",
        "Anomaly view: Color by 'anomaly_score' (continuous gradient).",
    ]:
        add_paragraph(doc, f"• {b}")

    OUTDIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))
    print(f"✅ DOCX created: {DOCX_PATH}")


if __name__ == "__main__":
    build_docx()
